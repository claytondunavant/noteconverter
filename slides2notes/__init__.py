from pdf2image import convert_from_bytes
from os import system
from os import path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Slides2Notes:

    numSlides = 0

    def slides2images(self):
        imgsDir = self.processID + "/imgs"
        system("mkdir " + imgsDir)

        images = convert_from_bytes(self.inputFile)

        for image in images:
            image.save(imgsDir + "/slide" + str(self.numSlides) + ".jpg", 'JPEG')
            self.numSlides = self.numSlides + 1

    def slides2docx(self):
        doc = Document()
        imgsDir = self.processID + "/imgs"

        for slideNum in range(self.numSlides):

            # add image of slide to the top center of the page
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("")
            run.add_picture(imgsDir + "/slide" + str(slideNum) + ".jpg", width=Inches(6))

            # insert page break at the end of the page
            doc.add_page_break()

        doc.save(self.processID + "/" + self.inputFileName + ".docx")


    def run(self):
        #clean
        system("rm -rf " + self.processID)
        # create file for current id
        system("mkdir " + self.processID)

        #turn the slides into images        
        self.slides2images()

        #create docx
        self.slides2docx()


        # delete current id file
        #system("rm -rf" + self.processID)

    def __init__(self, processID, inputFile, outputType):
        self.processID = str(processID)
        self.inputFile = inputFile
        self.inputFileName = path.splitext(inputFile.filename)[0]
        self.outputType = outputType
        self.run()

    
